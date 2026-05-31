"""Compiler Agent for UC-I-1 Hot Science candidate documents."""

from __future__ import annotations

import json
import csv
from collections import defaultdict
from dataclasses import dataclass, field
from html import escape
from pathlib import Path

from agents.hot_science.config import SourceConfig
from agents.hot_science.date_utils import parse_target_month, scan_window_for_month
from agents.hot_science.schema import CandidateRecord, ScoreItem


@dataclass(frozen=True)
class CompiledCandidateSet:
    target_month: str
    user_criteria: str | None
    candidates: list[CandidateRecord]
    excluded: list[CandidateRecord]
    manual_review: list[CandidateRecord]
    honorable_mentions: list[CandidateRecord]
    clusters: dict[str, list[CandidateRecord]]
    preprints: list[CandidateRecord] = field(default_factory=list)
    watchlist: list[CandidateRecord] = field(default_factory=list)
    source_errors: list[CandidateRecord] = field(default_factory=list)
    source_diagnostics: dict[str, dict[str, str | int]] = field(default_factory=dict)
    sources: list[SourceConfig] = field(default_factory=list)
    rubric_version: str | None = None
    standing_criteria_version: str = "hot_science_v2"

    def to_json_dict(self) -> dict:
        return {
            "target_month": self.target_month,
            "user_criteria": self.user_criteria,
            "run_config": {
                "target_month": self.target_month,
                "target_month_window": _target_month_window_label(self.target_month),
                "source_scan_window": _source_scan_window_label(self.target_month),
                "search_focus": _criteria_summary(self.user_criteria),
                "standing_criteria_version": self.standing_criteria_version,
                "rubric_version": self.rubric_version or _infer_rubric_version(self.candidates),
                "date_eligibility_rule": (
                    "Use the primary work's canonical online publication date; "
                    "press, posted, submitted, received, updated, and deposit dates do not qualify."
                ),
            },
            "retrieval_method": _retrieval_method_summary(self.user_criteria),
            "source_inventory": [_source_inventory_row(source) for source in self.sources],
            "counts": {
                "total_paper_records_categorized": _total_paper_records_categorized(self),
                "candidates": len(self.candidates),
                "preprints": len(self.preprints),
                "watchlist": len(self.watchlist),
                "excluded": len(self.excluded),
                "manual_review": len(self.manual_review),
                "honorable_mentions": len(self.honorable_mentions),
                "source_errors": len(self.source_errors),
                "clusters": len(self.clusters),
            },
            "zero_count_explanations": _zero_count_explanations(self),
            "candidates": [candidate.to_dict() for candidate in self.candidates],
            "clusters": {
                cluster: [candidate.to_dict() for candidate in candidates]
                for cluster, candidates in self.clusters.items()
            },
            "honorable_mentions": [candidate.to_dict() for candidate in self.honorable_mentions],
            "preprints": [candidate.to_dict() for candidate in self.preprints],
            "watchlist": [candidate.to_dict() for candidate in self.watchlist],
            "manual_review": [candidate.to_dict() for candidate in self.manual_review],
            "excluded": [candidate.to_dict() for candidate in self.excluded],
            "source_errors": [candidate.to_dict() for candidate in self.source_errors],
            "source_diagnostics": self.source_diagnostics,
        }


class CompilerAgent:
    """Cluster and assemble candidate sets into platform/doc-ready output."""

    def compile(
        self,
        *,
        target_month: str,
        candidates: list[CandidateRecord],
        excluded: list[CandidateRecord],
        manual_review: list[CandidateRecord] | None = None,
        preprints: list[CandidateRecord] | None = None,
        user_criteria: str | None = None,
        sources: list[SourceConfig] | tuple[SourceConfig, ...] | None = None,
        source_errors: list[CandidateRecord] | None = None,
        rubric_version: str | None = None,
        standing_criteria_version: str = "hot_science_v2",
    ) -> CompiledCandidateSet:
        clusters: dict[str, list[CandidateRecord]] = defaultdict(list)
        sorted_candidates = sorted(
            candidates,
            key=lambda c: c.significance.composite_score or 0,
            reverse=True,
        )
        for candidate in sorted_candidates:
            clusters[candidate.theme_cluster or "general_climate_research"].append(candidate)
            candidate.add_audit("compiler", "clustered", candidate.theme_cluster)

        honorable = [
            candidate for candidate in sorted_candidates if candidate.honorable_mention_candidate
        ]
        separated_preprints, remaining_excluded = _split_preprints(
            excluded,
            explicit_preprints=preprints or [],
        )
        watchlist, remaining_manual_review = _split_watchlist(manual_review or [])
        source_error_list = source_errors or []
        source_diagnostics = _source_diagnostic_rows(
            candidate_rows=sorted_candidates,
            preprint_rows=separated_preprints,
            watchlist_rows=watchlist,
            manual_review_rows=remaining_manual_review,
            excluded_rows=remaining_excluded,
            sources=sources,
            source_errors=source_error_list,
        )
        return CompiledCandidateSet(
            target_month=target_month,
            user_criteria=user_criteria,
            candidates=sorted_candidates,
            excluded=remaining_excluded,
            manual_review=remaining_manual_review,
            honorable_mentions=honorable,
            clusters=dict(clusters),
            preprints=separated_preprints,
            watchlist=watchlist,
            source_errors=source_error_list,
            source_diagnostics=source_diagnostics,
            sources=list(sources or []),
            rubric_version=rubric_version or _infer_rubric_version(sorted_candidates),
            standing_criteria_version=standing_criteria_version,
        )

    def write_json(self, compiled: CompiledCandidateSet, path: str | Path) -> None:
        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(json.dumps(compiled.to_json_dict(), indent=2, sort_keys=True))

    def write_markdown(self, compiled: CompiledCandidateSet, path: str | Path) -> None:
        """Write a Word-doc-friendly Markdown draft for review."""
        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        lines: list[str] = [
            f"# Hot Science Candidate Set — {compiled.target_month}",
            "",
            "## Run Configuration",
            "",
            f"- Search focus: {_criteria_summary(compiled.user_criteria)}",
            f"- Target month eligibility window: {compiled.target_month} ({_target_month_window_label(compiled.target_month)})",
            f"- Source retrieval scan window: {_source_scan_window_label(compiled.target_month)}",
            f"- Standing criteria version: {compiled.standing_criteria_version}",
            f"- Rubric version: {compiled.rubric_version or _infer_rubric_version(compiled.candidates)}",
            "- Date eligibility: primary work canonical online publication date only",
            "- Strict inclusion policy: top candidates must pass primary-source, date, direct-climate-lens, domain-fit, substantive-finding, and exclusion gates",
            f"- Total paper records categorized: {_total_paper_records_categorized(compiled)}",
            f"- Category breakdown: {_category_breakdown(compiled)}",
            f"- Top candidates: {len(compiled.candidates)}",
            f"- Preprints: {len(compiled.preprints)}",
            f"- Watchlist: {len(compiled.watchlist)}",
            f"- Manual-review queue: {len(compiled.manual_review)}",
            f"- Honorable mention candidates: {len(compiled.honorable_mentions)}",
            f"- Excluded candidates: {len(compiled.excluded)}",
            f"- Source errors: {len(compiled.source_errors)}",
            "",
        ]
        zero_explanations = _zero_count_explanations(compiled)
        if zero_explanations:
            lines.extend(["**Why any categories are zero**", ""])
            lines.extend(f"- {explanation}" for explanation in zero_explanations)
            lines.append("")
        lines.extend(format_search_methodology(compiled))
        lines.extend(format_source_inventory(compiled.sources))
        lines.append("## Top Candidates (Strict Include Only)")
        lines.append("")
        if compiled.candidates:
            for candidate in compiled.candidates:
                lines.extend(format_candidate(candidate, bucket="candidate"))
        else:
            lines.append("- No evaluated candidates.")
            lines.append("")
        if compiled.preprints:
            lines.append("## Preprints")
            lines.append("")
            for candidate in compiled.preprints:
                lines.extend(format_candidate(candidate, bucket="preprint"))
        if compiled.watchlist:
            lines.append("## Non-Target-Month Watchlist")
            lines.append("")
            for candidate in compiled.watchlist:
                lines.extend(format_candidate(candidate, bucket="watchlist"))
        if compiled.manual_review:
            lines.append("## Manual Review Queue")
            lines.append("")
            for candidate in compiled.manual_review:
                lines.extend(format_candidate(candidate, bucket="manual_review"))
        if compiled.honorable_mentions:
            lines.append("## Honorable Mentions")
            lines.append("")
            for candidate in compiled.honorable_mentions:
                lines.extend(format_candidate(candidate, bucket="honorable_mention"))
        if compiled.excluded:
            lines.append("## Excluded Candidates Appendix")
            lines.append("")
            for candidate in compiled.excluded:
                reasons = "; ".join(
                    f"{flag.code}: {flag.rationale}" for flag in candidate.exclusion_flags
                )
                lines.append(f"- {candidate.title} — {reasons or 'No reason recorded'}")
        if compiled.source_diagnostics:
            lines.append("")
            lines.append("## Source Diagnostics")
            lines.append("")
            lines.extend(format_source_diagnostics_table(compiled.source_diagnostics))
        lines.extend(format_criteria_brief(compiled.user_criteria))
        output_path.write_text("\n".join(lines) + "\n")

    def write_docx(self, compiled: CompiledCandidateSet, path: str | Path) -> None:
        """Write the primary team-facing Word document.

        This document intentionally includes top candidates and manual-review
        candidates, while excluding the debug appendices used for QA.
        """
        try:
            from docx import Document
            from docx.enum.table import WD_ALIGN_VERTICAL
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Inches, Pt, RGBColor
        except ImportError as exc:
            raise RuntimeError("python-docx is required to write Hot Science DOCX output.") from exc

        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        _configure_docx_styles(doc, Pt, RGBColor)

        title = doc.add_paragraph()
        title.style = doc.styles["Title"]
        title.add_run(f"Hot Science Candidate Set - {compiled.target_month}")
        subtitle = doc.add_paragraph()
        subtitle.style = doc.styles["Subtitle"]
        subtitle.add_run("Top candidates and manual-review candidates from the Hot Science agent run.")

        _add_docx_heading(doc, "At a Glance")
        _add_at_a_glance_docx(doc, compiled, OxmlElement, qn, WD_ALIGN_VERTICAL)

        _add_docx_heading(doc, "Run Configuration")
        _add_key_value_table_docx(
            doc,
            _run_configuration_rows(compiled),
            OxmlElement,
            qn,
            WD_ALIGN_VERTICAL,
        )

        zero_explanations = _zero_count_explanations(compiled)
        if zero_explanations:
            _add_docx_heading(doc, "Why Any Categories Are Zero", level=2)
            for explanation in zero_explanations:
                _add_docx_bullet(doc, explanation)

        _add_docx_heading(doc, "Search Criteria and Methodology")
        for item in _search_methodology_docx_items(compiled):
            _add_docx_bullet(doc, item)

        _add_docx_heading(doc, "Data Sources Searched or Configured")
        _add_source_inventory_docx(doc, compiled.sources, OxmlElement, qn, WD_ALIGN_VERTICAL)

        _add_docx_heading(doc, "Top Candidates")
        if compiled.candidates:
            for index, candidate in enumerate(compiled.candidates, start=1):
                _add_candidate_docx(
                    doc,
                    candidate,
                    bucket="candidate",
                    index=index,
                    OxmlElement=OxmlElement,
                    qn=qn,
                    WD_ALIGN_VERTICAL=WD_ALIGN_VERTICAL,
                )
        else:
            doc.add_paragraph("No papers passed every strict include gate.")

        _add_docx_heading(doc, "Manual Review Candidates")
        if compiled.manual_review:
            for index, candidate in enumerate(compiled.manual_review, start=1):
                _add_candidate_docx(
                    doc,
                    candidate,
                    bucket="manual_review",
                    index=index,
                    OxmlElement=OxmlElement,
                    qn=qn,
                    WD_ALIGN_VERTICAL=WD_ALIGN_VERTICAL,
                )
        else:
            doc.add_paragraph("No papers remained ambiguous after source, date, and fit checks.")

        if compiled.source_diagnostics:
            _add_docx_heading(doc, "Source Diagnostics")
            _add_source_diagnostics_docx(
                doc,
                compiled.source_diagnostics,
                OxmlElement,
                qn,
                WD_ALIGN_VERTICAL,
            )

        if compiled.user_criteria and _is_long_criteria(compiled.user_criteria):
            doc.add_page_break()
            _add_docx_heading(doc, "Full Criteria Brief")
            for line in compiled.user_criteria.splitlines():
                paragraph = doc.add_paragraph()
                if not line:
                    paragraph.add_run("")
                    continue
                paragraph.add_run(line)

        doc.save(output_path)

    def write_review_csv(self, compiled: CompiledCandidateSet, path: str | Path) -> None:
        """Write a spreadsheet-friendly review queue for the team voting workflow."""
        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        fields = [
            "bucket",
            "target_month",
            "user_criteria",
            "routing_reason",
            "watchlist_reason",
            "rubric_version",
            "theme_cluster",
            "title",
            "authors",
            "venue",
            "online_publication_date",
            "doi",
            "primary_url",
            "pdf_url",
            "open_access",
            "paywall",
            "abstract_accessible",
            "press_coverage",
            "prior_edition",
            "composite_score",
            "confidence",
            "novelty_subtype",
            "novelty_evidence",
            "impact_evidence",
            "eligibility_verdict",
            "fit_passed",
            "fit_manual_review_reason",
            "standing_scope_aligned",
            "run_focus_aligned",
            "date_eligible",
            "primary_object_verified",
            "fit_relevance_claim",
            "fit_evidence_source",
            "fit_evidence_snippet",
            "topic_tags",
            "exclusion_flags",
            "missing_reasons",
        ]
        rows = [
            *(
                _candidate_row("candidate", candidate, compiled.target_month, compiled.user_criteria)
                for candidate in compiled.candidates
            ),
            *(
                _candidate_row("preprint", candidate, compiled.target_month, compiled.user_criteria)
                for candidate in compiled.preprints
            ),
            *(
                _candidate_row("watchlist", candidate, compiled.target_month, compiled.user_criteria)
                for candidate in compiled.watchlist
            ),
            *(
                _candidate_row("manual_review", candidate, compiled.target_month, compiled.user_criteria)
                for candidate in compiled.manual_review
            ),
            *(
                _candidate_row("excluded", candidate, compiled.target_month, compiled.user_criteria)
                for candidate in compiled.excluded
            ),
        ]
        with output_path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.DictWriter(handle, fieldnames=fields)
            writer.writeheader()
            writer.writerows(rows)

    def write_source_breakdown_csv(
        self,
        compiled: CompiledCandidateSet,
        path: str | Path,
        sources: list[SourceConfig] | tuple[SourceConfig, ...] | None = None,
        source_errors: list[CandidateRecord] | None = None,
    ) -> None:
        """Write per-source bucket counts for coverage/routing diagnostics.

        One candidate can be attributed to more than one source if duplicates
        were consolidated. Within each bucket, a candidate is counted once per
        contributing source.
        """
        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        rows = _source_diagnostic_rows(
            candidate_rows=compiled.candidates,
            preprint_rows=compiled.preprints,
            watchlist_rows=compiled.watchlist,
            manual_review_rows=compiled.manual_review,
            excluded_rows=compiled.excluded,
            sources=sources,
            source_errors=source_errors or compiled.source_errors,
        )

        fields = [
            "source_id",
            "source_name",
            "source_type",
            "enabled",
            "verified",
            "preprints",
            "watchlist",
            "manual_review",
            "excluded",
            "source_errors",
            "total",
        ]
        configured_sources = list(sources or [])
        ordered_ids = [source.id for source in configured_sources if source.id in rows] + sorted(
            set(rows) - {source.id for source in configured_sources}
        )
        with output_path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.DictWriter(handle, fieldnames=fields)
            writer.writeheader()
            for source_id in ordered_ids:
                writer.writerow(rows[source_id])


def format_search_methodology(compiled: CompiledCandidateSet) -> list[str]:
    return [
        "## Search Criteria and Methodology",
        "",
        "**Search criteria used**",
        "",
        f"- User-provided focus: {_criteria_summary(compiled.user_criteria)}",
        f"- Criteria brief length: {len(compiled.user_criteria or '')} characters",
        f"- Target month eligibility filter: {compiled.target_month} ({_target_month_window_label(compiled.target_month)})",
        f"- Source retrieval scan window: {_source_scan_window_label(compiled.target_month)}",
        "- Standing scope: climate-relevant peer-reviewed research, attribution studies, and official institutional data releases.",
        "- Date rule: the primary work's canonical online publication date must fall inside the target month; press, posted, received, updated, submitted, and repository deposit dates do not qualify by themselves.",
        "",
        "**How retrieval worked**",
        "",
        "1. Source Monitor scanned enabled scholarly APIs, journal feeds, institutional feeds, attribution sources, and preprint metadata endpoints.",
        "2. API searches use the user focus first, then the configured broad Hot Science climate queries where the source supports query variants.",
        "3. Journal RSS sources with ISSNs are supplemented by OpenAlex and Crossref ISSN/month backfills so article-level records are not dependent on RSS timing alone.",
        "4. Records are deduplicated by DOI, source URL, and normalized title before verification.",
        "5. The resolver, verifier, access, evaluator, coverage, prior-edition, and compiler agents route each record into candidates, preprints, watchlist, manual review, or exclusions.",
        "",
    ]


def format_criteria_brief(user_criteria: str | None) -> list[str]:
    if not user_criteria or not _is_long_criteria(user_criteria):
        return []
    lines = [
        "",
        "## Full Criteria Brief",
        "",
        "The full criteria brief is shown as quoted text so its Markdown headings do not change the report structure.",
        "",
    ]
    for line in user_criteria.splitlines():
        lines.append(f"> {line}" if line else ">")
    lines.append("")
    return lines


def format_source_inventory(sources: list[SourceConfig]) -> list[str]:
    if not sources:
        return []
    lines = [
        "## Data Sources Searched or Configured",
        "",
        "Enabled sources are scanned during the run. Disabled sources are shown for transparency but are not queried.",
        "",
        "**Enabled sources scanned**",
        "",
    ]
    for source in sources:
        if not source.enabled:
            continue
        row = _source_inventory_row(source)
        lines.append(
            f"- {row['name']} ({_source_kind_label(row['kind'])}, {_source_type_label(row['source_type'])}): "
            f"{row['retrieval_strategy']} Endpoint/coverage: {row['endpoint']}"
        )
    disabled_sources = [source for source in sources if not source.enabled]
    if disabled_sources:
        lines.extend(["", "**Configured but not scanned**", ""])
        for source in disabled_sources:
            row = _source_inventory_row(source)
            lines.append(
                f"- {row['name']} ({_source_kind_label(row['kind'])}, "
                f"{_source_type_label(row['source_type'])}): {row['endpoint']}"
            )
    lines.extend([
        "",
        "| Source | Kind | Type | Status | Endpoint / coverage | Retrieval strategy |",
        "| --- | --- | --- | --- | --- | --- |",
    ])
    for source in sources:
        row = _source_inventory_row(source)
        display_values = [
            row["name"],
            _source_kind_label(row["kind"]),
            _source_type_label(row["source_type"]),
            row["status"],
            row["endpoint"],
            row["retrieval_strategy"],
        ]
        lines.append(
            "| "
            + " | ".join(_md_cell(str(value)) for value in display_values)
            + " |"
        )
    lines.append("")
    return lines


def format_source_diagnostics_table(
    source_diagnostics: dict[str, dict[str, str | int]]
) -> list[str]:
    lines = [
        "| Source | Enabled | Type | Candidates | Preprints | Watchlist | Manual review | Excluded | Source errors |",
        "| --- | --- | --- | ---: | ---: | ---: | ---: | ---: | ---: |",
    ]
    for source_id, row in source_diagnostics.items():
        lines.append(
            "| "
            + " | ".join(
                [
                    _md_cell(f"{row['source_name']} ({source_id})"),
                    _md_cell(str(row["enabled"])),
                    _md_cell(_source_type_label(str(row["source_type"]))),
                    str(row["verified"]),
                    str(row["preprints"]),
                    str(row["watchlist"]),
                    str(row["manual_review"]),
                    str(row["excluded"]),
                    str(row["source_errors"]),
                ]
            )
            + " |"
        )
    return [*lines, ""]


DOCX_TABLE_WIDTH_DXA = 9360
DOCX_TABLE_INDENT_DXA = 120
DOCX_CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
DOCX_HEADER_FILL = "E8EEF5"
DOCX_MUTED_FILL = "F2F4F7"


def _configure_docx_styles(doc, Pt, RGBColor) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 77, 120)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(89, 89, 89)
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size, color in (
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(18)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(7)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(5)


def _run_configuration_items(compiled: CompiledCandidateSet) -> list[str]:
    return [
        f"Search focus: {_criteria_summary(compiled.user_criteria)}",
        f"Target month eligibility window: {compiled.target_month} ({_target_month_window_label(compiled.target_month)})",
        f"Source retrieval scan window: {_source_scan_window_label(compiled.target_month)}",
        f"Standing criteria version: {compiled.standing_criteria_version}",
        f"Rubric version: {compiled.rubric_version or _infer_rubric_version(compiled.candidates)}",
        "Date eligibility: primary work canonical online publication date only",
        "Strict inclusion policy: top candidates must pass primary-source, date, direct-climate-lens, domain-fit, substantive-finding, and exclusion gates",
        f"Total paper records categorized: {_total_paper_records_categorized(compiled)}",
        f"Category breakdown: {_category_breakdown(compiled)}",
        f"Top candidates: {len(compiled.candidates)}",
        f"Manual-review candidates: {len(compiled.manual_review)}",
        f"Preprints: {len(compiled.preprints)}",
        f"Watchlist: {len(compiled.watchlist)}",
        f"Honorable mention candidates: {len(compiled.honorable_mentions)}",
        f"Excluded candidates: {len(compiled.excluded)}",
        f"Source errors: {len(compiled.source_errors)}",
    ]


def _run_configuration_rows(compiled: CompiledCandidateSet) -> list[tuple[str, str]]:
    return [
        ("Search focus", _criteria_summary(compiled.user_criteria)),
        (
            "Target month eligibility window",
            f"{compiled.target_month} ({_target_month_window_label(compiled.target_month)})",
        ),
        ("Source retrieval scan window", _source_scan_window_label(compiled.target_month)),
        ("Standing criteria version", compiled.standing_criteria_version),
        ("Rubric version", compiled.rubric_version or _infer_rubric_version(compiled.candidates) or "Not recorded"),
        ("Date eligibility", "Primary work canonical online publication date only"),
        (
            "Strict inclusion policy",
            "Top candidates must pass primary-source, date, direct-climate-lens, domain-fit, substantive-finding, and exclusion gates.",
        ),
    ]


def _add_at_a_glance_docx(doc, compiled, OxmlElement, qn, WD_ALIGN_VERTICAL) -> None:
    rows = [
        ("Total categorized", _total_paper_records_categorized(compiled), "All paper-like records assigned to an output bucket."),
        ("Top candidates", len(compiled.candidates), "Strict includes ready for team review."),
        ("Manual review", len(compiled.manual_review), "Plausible records that need human verification before final use."),
        ("Excluded", len(compiled.excluded), "Records removed by date, source, duplicate, fit, or exclusion checks."),
        ("Preprints", len(compiled.preprints), "Separated from peer-reviewed top candidates."),
        ("Watchlist", len(compiled.watchlist), "Relevant records retained outside the target-month final set."),
        ("Honorable mentions", len(compiled.honorable_mentions), "High-signal records that did not clear the top-candidate gates."),
        ("Source errors", len(compiled.source_errors), "Enabled sources that failed during scan."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _apply_table_geometry(
        table,
        [2100, 1260, 6000],
        OxmlElement,
        qn,
        WD_ALIGN_VERTICAL,
    )
    for idx, header in enumerate(("Category", "Count", "Meaning")):
        _set_docx_cell(table.rows[0].cells[idx], header, bold=True, fill=DOCX_HEADER_FILL, qn=qn, OxmlElement=OxmlElement)
    for label, count, meaning in rows:
        row = table.add_row().cells
        _set_docx_cell(row[0], label, qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[1], str(count), bold=True, qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[2], meaning, qn=qn, OxmlElement=OxmlElement)
    _apply_table_geometry(
        table,
        [2100, 1260, 6000],
        OxmlElement,
        qn,
        WD_ALIGN_VERTICAL,
    )


def _add_key_value_table_docx(
    doc,
    rows: list[tuple[str, str]],
    OxmlElement,
    qn,
    WD_ALIGN_VERTICAL,
) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row().cells
        _set_docx_cell(row[0], label, bold=True, fill=DOCX_MUTED_FILL, qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[1], value, qn=qn, OxmlElement=OxmlElement)
    _apply_table_geometry(table, [2700, 6660], OxmlElement, qn, WD_ALIGN_VERTICAL)


def _search_methodology_docx_items(compiled: CompiledCandidateSet) -> list[str]:
    return [
        f"User-provided focus: {_criteria_summary(compiled.user_criteria)}",
        f"Criteria brief length: {len(compiled.user_criteria or '')} characters",
        f"Target month eligibility filter: {compiled.target_month} ({_target_month_window_label(compiled.target_month)})",
        f"Source retrieval scan window: {_source_scan_window_label(compiled.target_month)}",
        "Standing scope: climate-relevant peer-reviewed research, attribution studies, and official institutional data releases.",
        "Date rule: the primary work's canonical online publication date must fall inside the target month; press, posted, received, updated, submitted, and repository deposit dates do not qualify by themselves.",
        "Retrieval scans enabled scholarly APIs, journal feeds, institutional feeds, attribution sources, and preprint metadata endpoints.",
        "Records are deduplicated by DOI, source URL, and normalized title before verification.",
    ]


def _add_docx_heading(doc, text: str, *, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_docx_bullet(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def _add_source_inventory_docx(doc, sources: list[SourceConfig], OxmlElement, qn, WD_ALIGN_VERTICAL) -> None:
    if not sources:
        doc.add_paragraph("No data sources were attached to this compiled candidate set.")
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Source", "Type", "Status", "Endpoint / coverage"]
    for index, header in enumerate(headers):
        _set_docx_cell(table.rows[0].cells[index], header, bold=True, fill=DOCX_HEADER_FILL, qn=qn, OxmlElement=OxmlElement)
    for source in sources:
        row = table.add_row().cells
        inventory = _source_inventory_row(source)
        _set_docx_cell(row[0], inventory["name"], qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(
            row[1],
            f"{_source_kind_label(inventory['kind'])} / {_source_type_label(inventory['source_type'])}",
            qn=qn,
            OxmlElement=OxmlElement,
        )
        _set_docx_cell(row[2], inventory["status"], qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[3], inventory["endpoint"], qn=qn, OxmlElement=OxmlElement)
    _apply_table_geometry(table, [2400, 2100, 1100, 3760], OxmlElement, qn, WD_ALIGN_VERTICAL)


def _add_source_diagnostics_docx(
    doc,
    source_diagnostics: dict[str, dict[str, str | int]],
    OxmlElement,
    qn,
    WD_ALIGN_VERTICAL,
) -> None:
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = [
        "Source",
        "On",
        "Type",
        "Top",
        "Manual",
        "Excl.",
        "Prep.",
        "Errors",
    ]
    for index, header in enumerate(headers):
        _set_docx_cell(table.rows[0].cells[index], header, bold=True, fill=DOCX_HEADER_FILL, qn=qn, OxmlElement=OxmlElement)
    for source_id, row_data in source_diagnostics.items():
        row = table.add_row().cells
        _set_docx_cell(row[0], str(row_data["source_name"]) or source_id, qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[1], str(row_data["enabled"]), qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[2], _source_type_label(str(row_data["source_type"])), qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[3], str(row_data["verified"]), qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[4], str(row_data["manual_review"]), qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[5], str(row_data["excluded"]), qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[6], str(row_data["preprints"]), qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[7], str(row_data["source_errors"]), qn=qn, OxmlElement=OxmlElement)
    _apply_table_geometry(table, [2700, 650, 1700, 700, 950, 850, 850, 960], OxmlElement, qn, WD_ALIGN_VERTICAL)


def _add_candidate_docx(
    doc,
    candidate: CandidateRecord,
    *,
    bucket: str,
    index: int,
    OxmlElement,
    qn,
    WD_ALIGN_VERTICAL,
) -> None:
    title = doc.add_heading(f"{index}. {candidate.title}", level=2)
    title.paragraph_format.keep_with_next = True
    rows = [
        ("Eligibility verdict", _eligibility_verdict(candidate, bucket)),
        ("Why this is here", _routing_reason_label(candidate, bucket)),
        ("Authors", ", ".join(candidate.authors[:8]) or "Unknown"),
        ("Journal/source", candidate.publication.venue or "Unknown"),
        ("Online publication date", candidate.publication.online_publication_date or "Unknown"),
        ("DOI", f"https://doi.org/{candidate.doi}" if candidate.doi else "DOI unresolved"),
        ("Primary URL", candidate.publication.primary_source_url or candidate.publication.url or "Unknown"),
        ("Access", _access_label(candidate)),
        ("Direct Hot Science fit gate", _team_fit_gate_label(candidate)),
        ("Fit evidence source", candidate.fit_assessment.evidence_source or "Manual review required"),
        ("Topic tags", ", ".join(candidate.topic_tags) or "None recorded"),
        ("Discovered via", _discovery_label(candidate)),
    ]
    if bucket == "manual_review":
        rows.append(
            (
                "Manual-review reason",
                candidate.fit_assessment.manual_review_reason
                or candidate.routing_reason
                or _missing_reason_summary(candidate)
                or "Needs human review before scoring.",
            )
        )
    if candidate.significance.composite_score is not None:
        rows.append(
            (
                "Composite score",
                f"{candidate.significance.composite_score} ({candidate.significance.overall_confidence} confidence)",
            )
        )

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row().cells
        _set_docx_cell(row[0], label, bold=True, fill=DOCX_MUTED_FILL, qn=qn, OxmlElement=OxmlElement)
        _set_docx_cell(row[1], value, qn=qn, OxmlElement=OxmlElement)
    _apply_table_geometry(table, [2100, 7260], OxmlElement, qn, WD_ALIGN_VERTICAL)

    doc.add_paragraph("Abstract").runs[0].bold = True
    doc.add_paragraph(_truncate(candidate.abstract or "Unavailable", 1200))

    if candidate.press_coverage:
        doc.add_paragraph("Popular press coverage").runs[0].bold = True
        for press in candidate.press_coverage:
            _add_docx_bullet(doc, f"{press.outlet}: {press.headline or press.url} - {press.url}")


def _set_docx_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    fill: str | None = None,
    qn=None,
    OxmlElement=None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    if fill and qn is not None and OxmlElement is not None:
        _shade_docx_cell(cell, fill, OxmlElement, qn)


def _apply_table_geometry(
    table,
    widths: list[int],
    OxmlElement,
    qn,
    WD_ALIGN_VERTICAL,
) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(DOCX_TABLE_INDENT_DXA))

    grid = tbl.tblGrid
    for idx, width in enumerate(widths):
        if idx >= len(grid.gridCol_lst):
            grid.append(OxmlElement("w:gridCol"))
        grid.gridCol_lst[idx].set(qn("w:w"), str(width))

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell, OxmlElement, qn)


def _set_cell_margins(cell, OxmlElement, qn) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, width in DOCX_CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def _shade_docx_cell(cell, fill: str, OxmlElement, qn) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _team_fit_gate_label(candidate: CandidateRecord) -> str:
    if candidate.fit_assessment.passed is True:
        return "Fit verified"
    if candidate.fit_assessment.passed is False:
        return "Fit failed"
    return candidate.fit_assessment.manual_review_reason or "Manual review required"


def format_candidate(candidate: CandidateRecord, *, bucket: str = "candidate") -> list[str]:
    sig = candidate.significance
    doi = f"https://doi.org/{candidate.doi}" if candidate.doi else "DOI unresolved"
    lines = [
        _paper_heading(candidate.title),
        "",
        f"- Bucket: {bucket.replace('_', ' ')}",
        f"- Eligibility verdict: {_eligibility_verdict(candidate, bucket)}",
        f"- Why this is here: {_routing_reason_label(candidate, bucket)}",
        f"- Authors: {', '.join(candidate.authors[:6]) or 'Unknown'}",
        f"- Journal/source: {candidate.publication.venue or 'Unknown'}",
        f"- Online publication date: {candidate.publication.online_publication_date or 'Unknown'}",
        f"- DOI: {doi}",
        f"- Primary URL: {candidate.publication.primary_source_url or candidate.publication.url or 'Unknown'}",
        f"- Access: {_access_label(candidate)}",
        f"- Abstract: {candidate.abstract or 'Unavailable'}",
        "",
        "**Fit and verification evidence**",
        "",
        f"- Fit claim: {candidate.fit_assessment.relevance_claim or 'Not recorded'}",
        f"- Fit evidence source: {candidate.fit_assessment.evidence_source or 'Not recorded'}",
        f"- Fit evidence snippet: {candidate.fit_assessment.evidence_snippet or 'Not recorded'}",
        f"- Topic tags: {', '.join(candidate.topic_tags) or 'None recorded'}",
        f"- Primary work type: {candidate.publication.primary_work_type or candidate.publication.venue_type or 'Unknown'}",
        f"- Date source field: {candidate.publication.date_source_field or 'Unknown'}",
        f"- Raw publication date: {candidate.publication.raw_publication_date or 'Unknown'}",
        f"- Discovered via: {_discovery_label(candidate)}",
        "",
        "**Strict criteria gates**",
        "",
        f"- Primary source gate: {_gate_label(candidate.verification.primary_object_verified)}",
        f"- Target-month date gate: {_gate_label(candidate.date_eligibility.eligible)}",
        f"- Direct Hot Science fit gate: {_fit_gate_label(candidate)}",
        f"- Run-specific focus gate: {_run_focus_gate_label(candidate)}",
        f"- Exclusion screen: {_exclusion_gate_label(candidate)}",
        "",
    ]
    if sig.composite_score is None:
        lines.extend(
            [
                "**Scoring status**",
                "",
                "- Not scored yet. This record needs manual review before rubric scoring because the primary paper, DOI, publication date, or source type still needs verification.",
                "",
            ]
        )
    else:
        lines.extend(
            [
                f"- Composite score: {sig.composite_score} ({sig.overall_confidence} confidence)",
                "",
                "**Rubric breakdown**",
                "",
                _format_score_item("Novelty", sig.novelty),
                _format_score_item("Impact magnitude", sig.impact_magnitude),
                _format_score_item("Earth-system signal", sig.earth_system_signal),
                _format_score_item("Cross-disciplinary additive signal", sig.cross_disciplinary),
                _format_score_item("Cascading-impact additive signal", sig.cascading_impact),
                "",
            ]
        )
    lines.extend(["**Popular press coverage**", ""])
    if candidate.press_coverage:
        for press in candidate.press_coverage:
            lines.append(f"- [AI-DISCOVERED PRESS] {press.outlet}: {press.headline or press.url} — {press.url}")
    else:
        lines.append("- None found")
    if candidate.prior_editions:
        lines.append("")
        lines.append("**Prior edition check**")
        lines.append("")
        for prior in candidate.prior_editions:
            lines.append(
                f"- Possible prior coverage: {prior.target_month or 'unknown month'} — {prior.title}"
            )
    if candidate.missing_reasons:
        lines.append("")
        lines.append("**Missing or manual-check fields**")
        lines.append("")
        for field, reason in candidate.missing_reasons.items():
            lines.append(f"- {field}: {reason}")
    lines.append("")
    return lines


def _paper_heading(title: str) -> str:
    """Return a slightly larger paper heading without promoting it to a section heading."""
    return f'### <span style="font-size: 1.08em;">{escape(title, quote=False)}</span>'


def _candidate_row(
    bucket: str,
    candidate: CandidateRecord,
    target_month: str,
    user_criteria: str | None,
) -> dict[str, str]:
    return {
        "bucket": bucket,
        "target_month": candidate.target_month or target_month,
        "user_criteria": user_criteria or "",
        "routing_reason": candidate.routing_reason or "",
        "watchlist_reason": candidate.watchlist_reason or "",
        "rubric_version": candidate.rubric_version or candidate.significance.rubric_version or "",
        "theme_cluster": candidate.theme_cluster or "",
        "title": candidate.title,
        "authors": "; ".join(candidate.authors),
        "venue": candidate.publication.venue or "",
        "online_publication_date": candidate.publication.online_publication_date or "",
        "doi": candidate.doi or "",
        "primary_url": candidate.publication.primary_source_url or candidate.publication.url or "",
        "pdf_url": candidate.publication.full_text_pdf_url or "",
        "open_access": _bool_label(candidate.publication.open_access),
        "paywall": _bool_label(candidate.publication.paywall),
        "abstract_accessible": _bool_label(candidate.publication.abstract_accessible),
        "press_coverage": "; ".join(press.url for press in candidate.press_coverage),
        "prior_edition": "; ".join(
            f"{prior.target_month or 'unknown'}: {prior.title}" for prior in candidate.prior_editions
        ),
        "composite_score": str(candidate.significance.composite_score or ""),
        "confidence": candidate.significance.overall_confidence,
        "novelty_subtype": candidate.significance.novelty.subtype or "",
        "novelty_evidence": candidate.significance.novelty.evidence or "",
        "impact_evidence": candidate.significance.impact_magnitude.evidence or "",
        "eligibility_verdict": _eligibility_verdict(candidate, bucket),
        "fit_passed": _bool_label(candidate.fit_assessment.passed),
        "fit_manual_review_reason": candidate.fit_assessment.manual_review_reason or "",
        "standing_scope_aligned": _bool_label(candidate.fit_assessment.standing_scope_aligned),
        "run_focus_aligned": _bool_label(candidate.fit_assessment.run_focus_aligned),
        "date_eligible": _bool_label(candidate.date_eligibility.eligible),
        "primary_object_verified": _bool_label(candidate.verification.primary_object_verified),
        "fit_relevance_claim": candidate.fit_assessment.relevance_claim or "",
        "fit_evidence_source": candidate.fit_assessment.evidence_source or "",
        "fit_evidence_snippet": candidate.fit_assessment.evidence_snippet or "",
        "topic_tags": "; ".join(candidate.topic_tags),
        "exclusion_flags": "; ".join(
            f"{flag.code}: {flag.rationale}" for flag in candidate.exclusion_flags
        ),
        "missing_reasons": "; ".join(
            f"{field}: {reason}" for field, reason in candidate.missing_reasons.items()
        ),
    }


def _target_month_window_label(target_month: str) -> str:
    try:
        window_start, window_end = parse_target_month(target_month)
    except ValueError:
        return "invalid target month"
    return f"{window_start.isoformat()} through {window_end.isoformat()}"


def _source_scan_window_label(target_month: str) -> str:
    try:
        window_start, window_end = scan_window_for_month(target_month)
    except ValueError:
        return "invalid target month"
    return f"{window_start.isoformat()} through {window_end.isoformat()}"


def _retrieval_method_summary(user_criteria: str | None) -> dict[str, str]:
    return {
        "user_focus": _criteria_summary(user_criteria),
        "query_strategy": (
            "Use the user focus as the lead query, then configured broad Hot Science climate "
            "queries for scholarly APIs and ISSN backfills where supported."
        ),
        "source_strategy": (
            "Scan enabled scholarly APIs, journal RSS feeds, institutional feeds, attribution "
            "sources, and preprint metadata endpoints."
        ),
        "date_filtering": (
            "Keep monthly candidates only when the primary work's canonical online publication "
            "date falls inside the target month."
        ),
        "deduplication": "Deduplicate by DOI first, then source URL and normalized title.",
        "routing": (
            "Route records into top candidates, preprints, non-target-month watchlist, manual "
            "review, excluded appendix, and source diagnostics."
        ),
    }


def _criteria_summary(user_criteria: str | None) -> str:
    if not user_criteria:
        return "Standard monthly Hot Science criteria"
    for line in user_criteria.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        lowered = stripped.casefold()
        if lowered.startswith("search query:"):
            return stripped
        if stripped.startswith("#"):
            continue
        return _truncate(stripped, 240)
    return _truncate(" ".join(user_criteria.split()), 240)


def _is_long_criteria(user_criteria: str) -> bool:
    return len(user_criteria) > 500 or "\n#" in user_criteria or user_criteria.lstrip().startswith("#")


def _truncate(value: str, max_length: int) -> str:
    if len(value) <= max_length:
        return value
    return value[: max_length - 1].rstrip() + "..."


def _source_inventory_row(source: SourceConfig) -> dict[str, str]:
    return {
        "id": source.id,
        "name": source.name,
        "kind": source.kind,
        "source_type": source.source_type,
        "status": "enabled" if source.enabled else "disabled",
        "endpoint": _source_endpoint(source),
        "issns": ", ".join(source.issns),
        "notes": source.notes or "",
        "retrieval_strategy": _source_retrieval_strategy(source),
    }


def _source_kind_label(kind: str) -> str:
    labels = {
        "api": "API",
        "rss": "RSS",
        "static": "Static source",
        "web": "Web page",
    }
    return labels.get(kind, kind.replace("_", " ").title() if kind else "Unknown")


def _source_type_label(source_type: str) -> str:
    labels = {
        "attribution": "Attribution report",
        "data_release": "Data release",
        "institutional_report": "Institutional report",
        "peer_reviewed_journal": "Peer-reviewed journal",
        "popular_press": "Popular press",
        "preprint": "Preprint",
        "scholarly_api": "Scholarly API",
    }
    return labels.get(
        source_type,
        source_type.replace("_", " ").title() if source_type else "Unattributed",
    )


def _source_endpoint(source: SourceConfig) -> str:
    if source.url:
        endpoint = source.url
    elif source.id == "openalex":
        endpoint = "https://api.openalex.org/works"
    elif source.id == "crossref":
        endpoint = "https://api.crossref.org/works"
    elif source.id == "semantic_scholar":
        endpoint = "https://api.semanticscholar.org/graph/v1/paper/search"
    else:
        endpoint = "No public endpoint configured"
    if source.issns:
        endpoint = f"{endpoint}; ISSNs: {', '.join(source.issns)}"
    if source.notes:
        endpoint = f"{endpoint}; note: {source.notes}"
    return endpoint


def _source_retrieval_strategy(source: SourceConfig) -> str:
    if not source.enabled:
        return "Configured for transparency, but not scanned in this run."
    if source.id == "openalex":
        return "Scholarly API search over works, filtered by target-month publication date."
    if source.id == "crossref":
        return "Scholarly DOI metadata search over journal articles, filtered by publication date."
    if source.id == "semantic_scholar":
        return "Scholarly API search when enabled with an API key."
    if source.kind == "oai_pmh":
        return "OAI-PMH metadata scan, parsed for title, authors, DOI, abstract, and date."
    if source.kind == "institutional_feed":
        return "Institutional feed scan, then target-month and fit verification."
    if source.kind == "rss" and source.source_type == "peer_reviewed_journal" and source.issns:
        return "Journal RSS plus OpenAlex/Crossref ISSN-month backfill."
    if source.kind == "rss":
        return "RSS feed scan, then target-month and fit verification."
    if source.kind == "preprint_feed":
        return "Preprint feed scan, then separate preprint bucket routing."
    if source.kind == "press_api":
        return "Press API scan when credentials or licensed access are configured."
    return "Configured source scan through the source monitor."


def _md_cell(value: str) -> str:
    return " ".join(value.replace("|", "\\|").split())


def _total_paper_records_categorized(compiled: CompiledCandidateSet) -> int:
    """Count paper-like records assigned to mutually exclusive output buckets."""
    return (
        len(compiled.candidates)
        + len(compiled.preprints)
        + len(compiled.watchlist)
        + len(compiled.manual_review)
        + len(compiled.excluded)
    )


def _category_breakdown(compiled: CompiledCandidateSet) -> str:
    return "; ".join(
        [
            f"top candidates {len(compiled.candidates)}",
            f"manual review {len(compiled.manual_review)}",
            f"excluded {len(compiled.excluded)}",
            f"preprints {len(compiled.preprints)}",
            f"watchlist {len(compiled.watchlist)}",
            f"honorable mentions {len(compiled.honorable_mentions)}",
            f"source errors {len(compiled.source_errors)}",
        ]
    )


def _zero_count_explanations(compiled: CompiledCandidateSet) -> list[str]:
    bucket_counts = [
        ("Top candidates", len(compiled.candidates), "No papers passed every strict include gate."),
        ("Preprints", len(compiled.preprints), "No retrieved records were classified as preprints."),
        (
            "Watchlist",
            len(compiled.watchlist),
            "No relevant wrong-month records were retained for the non-target-month watchlist.",
        ),
        (
            "Manual review",
            len(compiled.manual_review),
            "No papers remained ambiguous after source, date, and fit checks.",
        ),
        (
            "Honorable mention candidates",
            len(compiled.honorable_mentions),
            "No paper met the special high-audience-signal/lower-selection-score rule.",
        ),
        (
            "Excluded candidates",
            len(compiled.excluded),
            "No records failed source, date, duplicate, fit, or exclusion checks.",
        ),
        (
            "Source errors",
            len(compiled.source_errors),
            "All enabled sources completed without a scan failure.",
        ),
    ]
    return [
        f"{label}: {reason}"
        for label, count, reason in bucket_counts
        if count == 0
    ]


def _discovery_label(candidate: CandidateRecord) -> str:
    if not candidate.discovered_via:
        return "Not recorded"
    parts: list[str] = []
    for mention in candidate.discovered_via:
        label = mention.source or "Unknown source"
        if mention.source_type:
            label = f"{label} ({_source_type_label(mention.source_type)})"
        if mention.url:
            label = f"{label}: {mention.url}"
        parts.append(label)
    return "; ".join(parts)


def _candidate_source_ids(
    candidate: CandidateRecord,
    source_name_to_id: dict[str, str],
) -> set[str]:
    source_ids: set[str] = set()
    for mention in candidate.discovered_via:
        if mention.source in source_name_to_id:
            source_ids.add(source_name_to_id[mention.source])
        elif not source_name_to_id and mention.source:
            source_ids.add(mention.source)
    return source_ids


def _split_preprints(
    excluded: list[CandidateRecord],
    *,
    explicit_preprints: list[CandidateRecord],
) -> tuple[list[CandidateRecord], list[CandidateRecord]]:
    preprints = list(explicit_preprints)
    remaining: list[CandidateRecord] = []
    seen_ids = {candidate.candidate_id for candidate in preprints}
    for candidate in excluded:
        if _is_preprint_bucket(candidate):
            if candidate.candidate_id not in seen_ids:
                candidate.source_status = "preprint"
                candidate.routing_reason = candidate.routing_reason or "preprint_separate_bucket"
                preprints.append(candidate)
                seen_ids.add(candidate.candidate_id)
            continue
        remaining.append(candidate)
    return _sort_supporting_bucket(preprints), remaining


def _split_watchlist(
    manual_review: list[CandidateRecord],
) -> tuple[list[CandidateRecord], list[CandidateRecord]]:
    watchlist: list[CandidateRecord] = []
    remaining: list[CandidateRecord] = []
    for candidate in manual_review:
        if candidate.watchlist_reason or candidate.routing_reason == "non_target_month_watchlist":
            watchlist.append(candidate)
        else:
            remaining.append(candidate)
    return _sort_supporting_bucket(watchlist), remaining


def _is_preprint_bucket(candidate: CandidateRecord) -> bool:
    return (
        candidate.publication.venue_type == "preprint"
        or candidate.publication.primary_work_type == "preprint"
        or candidate.routing_reason == "preprint_separate_bucket"
        or any(flag.code == "preprint_bucket" for flag in candidate.exclusion_flags)
    )


def _sort_supporting_bucket(candidates: list[CandidateRecord]) -> list[CandidateRecord]:
    return sorted(
        candidates,
        key=lambda candidate: (
            candidate.publication.online_publication_date or "",
            candidate.title.casefold(),
        ),
        reverse=True,
    )


def _source_diagnostic_rows(
    *,
    candidate_rows: list[CandidateRecord],
    preprint_rows: list[CandidateRecord],
    watchlist_rows: list[CandidateRecord],
    manual_review_rows: list[CandidateRecord],
    excluded_rows: list[CandidateRecord],
    sources: list[SourceConfig] | tuple[SourceConfig, ...] | None = None,
    source_errors: list[CandidateRecord] | None = None,
) -> dict[str, dict[str, str | int]]:
    configured_sources = list(sources or [])
    rows: dict[str, dict[str, str | int]] = {}
    source_name_to_id = {source.name: source.id for source in configured_sources}

    for source in configured_sources:
        rows[source.id] = _empty_source_row(
            source.id,
            source.name,
            source.source_type,
            "yes" if source.enabled else "no",
        )

    for source_error in source_errors or []:
        for source_id in _candidate_source_ids(source_error, source_name_to_id) or {"unattributed"}:
            _ensure_source_row(rows, source_id)
            rows[source_id]["source_errors"] = int(rows[source_id]["source_errors"]) + 1

    for bucket, candidates in (
        ("verified", candidate_rows),
        ("preprints", preprint_rows),
        ("watchlist", watchlist_rows),
        ("manual_review", manual_review_rows),
        ("excluded", excluded_rows),
    ):
        for candidate in candidates:
            source_ids = _candidate_source_ids(candidate, source_name_to_id) or {"unattributed"}
            for source_id in source_ids:
                _ensure_source_row(rows, source_id)
                rows[source_id][bucket] = int(rows[source_id][bucket]) + 1
                rows[source_id]["total"] = int(rows[source_id]["total"]) + 1
    return rows


def _empty_source_row(
    source_id: str,
    source_name: str,
    source_type: str,
    enabled: str,
) -> dict[str, str | int]:
    return {
        "source_id": source_id,
        "source_name": source_name,
        "source_type": source_type,
        "enabled": enabled,
        "verified": 0,
        "preprints": 0,
        "watchlist": 0,
        "manual_review": 0,
        "excluded": 0,
        "source_errors": 0,
        "total": 0,
    }


def _ensure_source_row(rows: dict[str, dict[str, str | int]], source_id: str) -> None:
    if source_id not in rows:
        rows[source_id] = _empty_source_row(source_id, source_id, "", "")


def _routing_reason_label(candidate: CandidateRecord, bucket: str) -> str:
    if bucket == "candidate":
        return (
            candidate.fit_assessment.relevance_claim
            or candidate.routing_reason
            or "Evidence-backed Hot Science candidate for this target month."
        )
    if bucket == "preprint":
        return "Preprint separated from peer-reviewed primary candidates."
    if bucket == "watchlist":
        return candidate.watchlist_reason or "Relevant item outside the target month."
    if bucket == "manual_review":
        return (
            candidate.fit_assessment.manual_review_reason
            or candidate.routing_reason
            or _missing_reason_summary(candidate)
            or "Needs human review before scoring."
        )
    if bucket == "excluded":
        return _exclusion_summary(candidate) or candidate.routing_reason or "Excluded from candidate set."
    return candidate.routing_reason or "Included for reviewer context."


def _missing_reason_summary(candidate: CandidateRecord) -> str | None:
    if not candidate.missing_reasons:
        return None
    field, reason = next(iter(candidate.missing_reasons.items()))
    return f"{field}: {reason}"


def _exclusion_summary(candidate: CandidateRecord) -> str | None:
    if not candidate.exclusion_flags:
        return None
    return "; ".join(f"{flag.code}: {flag.rationale}" for flag in candidate.exclusion_flags)


def _infer_rubric_version(candidates: list[CandidateRecord]) -> str | None:
    for candidate in candidates:
        if candidate.rubric_version:
            return candidate.rubric_version
        if candidate.significance.rubric_version:
            return candidate.significance.rubric_version
    return None


def _format_score_item(label: str, item: ScoreItem) -> str:
    weight = f", weight {item.weight:g}" if item.weight is not None else ""
    subtype = f", {item.subtype}" if item.subtype else ""
    evidence = f" Evidence: {item.evidence}" if item.evidence else ""
    return f"- {label}: {item.score}{weight}{subtype} — {item.rationale}{evidence}"


def _eligibility_verdict(candidate: CandidateRecord, bucket: str) -> str:
    if bucket == "candidate" and candidate.fit_assessment.passed is True:
        return "Include"
    if bucket in {"manual_review", "watchlist", "preprint"}:
        return "Manual review"
    if bucket == "excluded" or candidate.exclusion_flags:
        return "Exclude"
    if candidate.fit_assessment.passed is False:
        return "Exclude"
    if candidate.fit_assessment.passed is None:
        return "Manual review"
    return "Unknown"


def _gate_label(value: bool | None) -> str:
    if value is True:
        return "Pass"
    if value is False:
        return "Fail"
    return "Not recorded"


def _fit_gate_label(candidate: CandidateRecord) -> str:
    if candidate.fit_assessment.passed is True:
        return "Pass"
    if candidate.fit_assessment.passed is False:
        return "Fail"
    return candidate.fit_assessment.manual_review_reason or "Manual review required"


def _run_focus_gate_label(candidate: CandidateRecord) -> str:
    if candidate.fit_assessment.run_focus_aligned is True:
        return "Pass"
    if candidate.fit_assessment.run_focus_aligned is False:
        return "Fail"
    return "No run-specific focus gate applied"


def _exclusion_gate_label(candidate: CandidateRecord) -> str:
    if candidate.exclusion_flags:
        return "Fail: " + _exclusion_summary(candidate)
    return "Pass"


def _access_label(candidate: CandidateRecord) -> str:
    publication = candidate.publication
    if publication.full_text_pdf_url:
        return f"Open PDF found ({publication.full_text_pdf_url})"
    if publication.paywall is True:
        return "Likely paywalled; abstract/metadata retained for review"
    if publication.paywall is False:
        return "Open or not paywalled according to source metadata"
    return publication.access_note or "Access status unknown"


def _bool_label(value: bool | None) -> str:
    if value is True:
        return "yes"
    if value is False:
        return "no"
    return "unknown"
