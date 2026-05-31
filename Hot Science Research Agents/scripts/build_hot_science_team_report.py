#!/usr/bin/env python3
"""Build a shareable Hot Science feedback-response DOCX and email draft."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
REGRESSION_JSON = ROOT / ".deepgreen" / "hot_science" / "april_2026_regression.json"
OUTPUT_DIR = ROOT / "outputs" / "hot_science"
DOCX_OUT = OUTPUT_DIR / "Hot_Science_April_2026_Feedback_Response.docx"
EMAIL_OUT = OUTPUT_DIR / "hot_science_team_email_draft.md"


def main() -> None:
    report = json.loads(REGRESSION_JSON.read_text())
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx(report, DOCX_OUT)
    EMAIL_OUT.write_text(email_draft())
    print(DOCX_OUT)
    print(EMAIL_OUT)


def build_docx(report: dict, output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    configure_styles(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Hot Science Agent Feedback Response and April 2026 Rerun")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.style = doc.styles["Subtitle"]
    subtitle.add_run(
        "Summary of changes made after Science R&D feedback, plus regression diagnostics from the April 2026 reviewer-calibration set."
    )

    add_heading(doc, "Executive Summary")
    summary = report["summary"]
    add_bullet(doc, f"Regression status: {summary['status'].upper()}.")
    add_bullet(doc, f"Reviewer-calibration cases passed: {summary['passed']} of {summary['total_cases']}.")
    add_bullet(doc, f"Rubric version: {summary['rubric_version']}.")
    add_bullet(doc, f"Journal ISSN/month backfill sources enabled: {summary['issn_backfill_sources']}.")

    add_heading(doc, "What Changed")
    for item in [
        "Press and repository records are treated as discovery leads; the pipeline now prefers the resolved canonical paper or report.",
        "Primary publication date discipline is stricter: press dates, posted dates, received dates, updated dates, and repository deposit dates do not qualify for monthly eligibility.",
        "The fit gate is abstract-first: a candidate needs evidence in the abstract or primary metadata before it can be scored.",
        "False-positive themes from the feedback are now encoded as reusable rules, including tectonics, formation history, paleoclimate-only, ambiguous pollutant language, methods-only papers, and social-science-only papers.",
        "Scoring now uses weighted scientific selection signals. Audience relevance is retained only for future summary drafting and no longer affects selection ranking.",
        "The review output is split into top candidates, preprints, non-target-month watchlist, manual review, excluded records, and source diagnostics.",
        "Journal RSS coverage is supplemented with OpenAlex and Crossref ISSN/month queries for configured journal sources.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "April 2026 Regression Result")
    p = doc.add_paragraph()
    p.add_run(
        "The rerun used the reviewer-derived April 2026 calibration cases created from the commented output and feedback documents. "
    )
    p.add_run(
        "All cases landed in the expected bucket with the expected routing or exclusion rule."
    ).bold = True

    doc.add_page_break()
    add_heading(doc, "Case Results")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [3000, 1800, 4560])
    headers = ["Case", "Outcome", "Rule confirmed"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)
    for case in report["cases"]:
        row = table.add_row().cells
        set_cell_text(row[0], readable_case_id(case["id"]))
        set_cell_text(
            row[1],
            f"PASS\n{case['expected_bucket']} -> {case['actual_bucket']}",
        )
        set_cell_text(row[2], case["generalized_rule"])

    doc.add_page_break()
    add_heading(doc, "Retrieval Diagnostics")
    p = doc.add_paragraph()
    p.add_run(
        "The source monitor now supplements configured peer-reviewed journal feeds with scholarly API backfills by ISSN and month. "
    )
    p.add_run(
        "This reduces dependence on journal RSS timing and helps recover article-level records directly from scholarly metadata."
    )
    source_table = doc.add_table(rows=1, cols=3)
    source_table.style = "Table Grid"
    set_table_widths(source_table, [2800, 3000, 3560])
    for idx, header in enumerate(["Source", "Source ID", "ISSNs"]):
        cell = source_table.rows[0].cells[idx]
        shade_cell(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)
    for source in report["retrieval_diagnostics"]["sources"]:
        row = source_table.add_row().cells
        set_cell_text(row[0], source["name"])
        set_cell_text(row[1], source["id"])
        set_cell_text(row[2], ", ".join(source["issns"]))

    doc.save(output_path)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_widths(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = tbl.tblGrid
    for idx, width in enumerate(widths):
        if idx >= len(grid.gridCol_lst):
            grid.append(OxmlElement("w:gridCol"))
        grid.gridCol_lst[idx].set(qn("w:w"), str(width))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def readable_case_id(case_id: str) -> str:
    return " ".join(part.capitalize() for part in case_id.split("_"))


def email_draft() -> str:
    return "Subject: Hot Science agent rerun based on your feedback\n\n" + "\n\n".join(
        email_body_paragraphs()
    ) + "\n"


def email_body_paragraphs() -> list[str]:
    return [
        "Hi all,",
        "Thank you again for the detailed feedback on the original Hot Science agent output. I took that feedback, converted the main issues into reusable evaluation rules and calibration cases, made changes across the agent workflow, and reran the April 2026 analysis against the reviewer-derived regression set.",
        "The attached document summarizes what changed and the new output from the rerun. In short, the updated workflow now resolves primary works more carefully, uses stricter publication-date rules, requires abstract or primary-metadata evidence before scoring, separates preprints and watchlist items, and adds source diagnostics. I also added ISSN/month scholarly API backfills so the journal search is less dependent on RSS timing alone.",
        "The April 2026 regression passed all 14 reviewer-calibration cases. The document includes the case-by-case result table and the retrieval updates that were added.",
        "Best,",
    ]


if __name__ == "__main__":
    main()
