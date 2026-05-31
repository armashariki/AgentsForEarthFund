#!/usr/bin/env python3
"""Build a Word version of the Hot Science criteria Markdown file."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DEFAULT_INPUT = ROOT / "outputs" / "hot_science" / "april_2026_climate_change_global_warming_criteria.md"
DEFAULT_OUTPUT = ROOT / "outputs" / "hot_science" / "Hot_Science_April_2026_Agent_Criteria.docx"


def main() -> None:
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_INPUT
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUTPUT
    criteria_text = input_path.read_text(encoding="utf-8")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)
    add_front_matter(document, input_path)
    add_markdown_content(document, criteria_text)
    add_footer(document)
    document.save(output_path)
    print(output_path)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor.from_string("555555")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)

    heading_specs = [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]
    for style_name, size, color, before, after in heading_specs:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Calibri"
    list_bullet.font.size = Pt(11)
    list_bullet.paragraph_format.left_indent = Inches(0.375)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.188)
    list_bullet.paragraph_format.space_after = Pt(4)
    list_bullet.paragraph_format.line_spacing = 1.25
    list_bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_front_matter(document: Document, input_path: Path) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("Hot Science Agent Criteria")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Climate Change and Global Warming Focus, April 2026")
    subtitle_run.italic = True
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string("555555")

    source = document.add_paragraph()
    source.paragraph_format.space_after = Pt(12)
    run = source.add_run(f"Source criteria file: {input_path}")
    run.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")


def add_markdown_content(document: Document, text: str) -> None:
    lines = text.splitlines()
    pending_paragraph: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line:
            flush_paragraph(document, pending_paragraph)
            continue
        if line.startswith("# "):
            flush_paragraph(document, pending_paragraph)
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            flush_paragraph(document, pending_paragraph)
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            flush_paragraph(document, pending_paragraph)
            document.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            flush_paragraph(document, pending_paragraph)
            para = document.add_paragraph(style="List Bullet")
            para.add_run(line[2:].strip())
            continue
        pending_paragraph.append(line.strip())

    flush_paragraph(document, pending_paragraph)


def flush_paragraph(document: Document, pending: list[str]) -> None:
    if not pending:
        return
    text = " ".join(pending)
    para = document.add_paragraph()
    if text.startswith("Search query:"):
        para.style = document.styles["Normal"]
        add_shading(para, "F4F6F9")
        run = para.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("0B2545")
    else:
        para.add_run(text)
    pending.clear()


def add_shading(paragraph, fill: str) -> None:
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(10)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run("Hot Science criteria reference")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("666666")
    paragraph.add_run("  |  Page ")
    add_page_number(paragraph)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    paragraph.add_run().add_break(WD_BREAK.LINE)


if __name__ == "__main__":
    main()
