"""Phase 1 tests for the Hot Science web service boundary."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from agents.hot_science.compiler import CompilerAgent
from agents.hot_science.config import SourceConfig
from agents.hot_science.evaluator import SignificanceEvaluatorAgent
from agents.hot_science.orchestrator import HotScienceRunResult
from agents.hot_science.schema import CandidateRecord, PublicationInfo, SourceMention
from web.backend.artifacts import HotScienceArtifactStore, primary_artifact
from web.backend.hot_science_service import HotScienceRunService
from web.backend.schemas import HotScienceRunRequest


def make_web_candidate(
    *,
    title: str = "Ocean warming accelerates Antarctic ice shelf retreat",
    doi: str | None = "10.1234/web.2026.001",
    date: str | None = "2026-04-12",
    venue_type: str = "peer_reviewed_journal",
) -> CandidateRecord:
    return CandidateRecord(
        title=title,
        doi=doi,
        authors=["Ada Ice", "Sam Ocean"],
        publication=PublicationInfo(
            venue="Nature Climate Change",
            venue_type=venue_type,
            online_publication_date=date,
            url=f"https://doi.org/{doi}" if doi else "https://example.org/no-doi",
            open_access=True,
        ),
        abstract=(
            "This study reports climate-change-driven ocean warming and Antarctic "
            "ice shelf retreat with implications for sea-level rise."
        ),
        abstract_source="publisher",
        discovered_via=[
            SourceMention(
                source="Nature Climate Change",
                url=f"https://doi.org/{doi}" if doi else "https://example.org/no-doi",
                date_seen=date,
                source_type=venue_type,
            )
        ],
        target_month="2026-04",
    )


def test_compiler_writes_team_docx_with_top_and_manual_review_only(tmp_path):
    top_candidate = SignificanceEvaluatorAgent().evaluate([make_web_candidate()]).evaluated[0]
    manual_review = make_web_candidate(
        title="Plausible climate impact paper needing abstract verification",
        doi="10.1234/manual.2026.001",
    )
    manual_review.abstract = None
    manual_review.routing_reason = "title_only_abstract_missing"
    manual_review.fit_assessment.manual_review_reason = "Needs abstract before scoring."
    excluded = make_web_candidate(
        title="Off-scope excluded paper title should not appear",
        doi="10.1234/excluded.2026.001",
    )
    excluded.add_exclusion("no_evidence_of_hot_science_fit", "No climate mechanism.")
    source = SourceConfig(
        id="nature_climate_change",
        name="Nature Climate Change",
        kind="rss",
        source_type="peer_reviewed_journal",
    )
    compiled = CompilerAgent().compile(
        target_month="2026-04",
        candidates=[top_candidate],
        excluded=[excluded],
        manual_review=[manual_review],
        user_criteria="Search query: climate change sea level rise",
        sources=(source,),
        rubric_version="hot_science_v2_candidate",
    )

    output = tmp_path / "hot_science.docx"
    CompilerAgent().write_docx(compiled, output)

    text = _docx_text(output)
    document = Document(output)
    assert "At a Glance" in text
    assert "Total categorized" in text
    assert "Run Configuration" in text
    assert "Top Candidates" in text
    assert "Manual Review Candidates" in text
    assert "Manual-review reason" in text
    assert "Ocean warming accelerates Antarctic ice shelf retreat" in text
    assert "Plausible climate impact paper needing abstract verification" in text
    assert "Off-scope excluded paper title should not appear" not in text
    assert "Excluded Candidates Appendix" not in text
    assert "Fit: Not recorded" not in text
    assert "Data Sources Searched or Configured" in text
    assert "Source Diagnostics" in text
    assert document.sections[0].top_margin.inches == 1
    assert document.sections[0].left_margin.inches == 1
    _assert_docx_tables_have_fixed_geometry(document)


def test_artifact_store_writes_docx_and_manifest(tmp_path):
    compiled = _compiled_fixture()
    result = HotScienceRunResult(
        run_id="run-123",
        target_month="2026-04",
        user_criteria="Search query: climate change sea level rise",
        raw_candidates=[],
        source_errors=[],
        compiled=compiled,
    )
    request = HotScienceRunRequest(
        target_month="2026-04",
        criteria_text="Search query: climate change sea level rise",
        requested_by="user1",
    ).normalized()
    store = HotScienceArtifactStore(tmp_path)

    artifacts = store.write_run_artifacts(result, request)
    primary = primary_artifact(artifacts)
    manifest = json.loads((tmp_path / "runs" / "run-123" / "manifest.json").read_text())

    assert primary.kind == "docx"
    assert Path(primary.path).exists()
    assert manifest["request"]["criteria_text"] == "Search query: climate change sea level rise"
    assert "criteria_file" not in manifest["request"]
    assert any(artifact.kind == "review_csv" for artifact in artifacts)
    try:
        store.resolve_artifact_path("run-123", "../escape.docx")
    except ValueError:
        pass
    else:
        raise AssertionError("Expected artifact path traversal to fail")


def test_hot_science_service_normalizes_request_and_returns_docx_link(tmp_path):
    class FakeRunner:
        captured: dict = {}

        def run(self, target_month, **kwargs):
            self.captured = {"target_month": target_month, **kwargs}
            return HotScienceRunResult(
                run_id="run-456",
                target_month=target_month,
                user_criteria=kwargs["user_criteria"],
                raw_candidates=[],
                source_errors=[],
                compiled=_compiled_fixture(user_criteria=kwargs["user_criteria"]),
            )

    runner = FakeRunner()
    artifact_store = HotScienceArtifactStore(tmp_path)
    service = HotScienceRunService(
        artifact_store=artifact_store,
        db_path=tmp_path / "candidates.sqlite3",
        runner=runner,
    )

    result = service.run(
        HotScienceRunRequest(
            target_month="2026-04",
            criteria_text="  Search query: climate change sea level rise  \n\nInclude April papers.  ",
            source_ids=("openalex", "crossref"),
            max_results_per_source=5,
            requested_by=" user1 ",
        )
    )

    assert runner.captured["target_month"] == "2026-04"
    assert runner.captured["retrieval_query"] == "climate change sea level rise"
    assert runner.captured["user_criteria"].endswith("Include April papers.")
    assert runner.captured["source_ids"] == {"openalex", "crossref"}
    assert result.primary_artifact.kind == "docx"
    assert result.primary_artifact.download_url.endswith(result.primary_artifact.filename)


def test_hot_science_service_empty_source_list_runs_all_enabled_sources(tmp_path):
    class FakeRunner:
        captured: dict = {}

        def run(self, target_month, **kwargs):
            self.captured = {"target_month": target_month, **kwargs}
            return HotScienceRunResult(
                run_id="run-789",
                target_month=target_month,
                user_criteria=kwargs["user_criteria"],
                raw_candidates=[],
                source_errors=[],
                compiled=_compiled_fixture(user_criteria=kwargs["user_criteria"]),
            )

    runner = FakeRunner()
    artifact_store = HotScienceArtifactStore(tmp_path)
    service = HotScienceRunService(
        artifact_store=artifact_store,
        db_path=tmp_path / "candidates.sqlite3",
        runner=runner,
    )

    service.run(
        HotScienceRunRequest(
            target_month="2026-04",
            criteria_text="Search query: climate change sea level rise",
            source_ids=("", "   "),
        )
    )

    assert runner.captured["source_ids"] is None


def _compiled_fixture(user_criteria: str = "Search query: climate change sea level rise"):
    candidate = SignificanceEvaluatorAgent().evaluate([make_web_candidate()]).evaluated[0]
    source = SourceConfig(
        id="nature_climate_change",
        name="Nature Climate Change",
        kind="rss",
        source_type="peer_reviewed_journal",
    )
    return CompilerAgent().compile(
        target_month="2026-04",
        candidates=[candidate],
        excluded=[],
        manual_review=[],
        user_criteria=user_criteria,
        sources=(source,),
        rubric_version="hot_science_v2_candidate",
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _assert_docx_tables_have_fixed_geometry(document: Document) -> None:
    assert document.tables
    for table in document.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        assert tbl_w is not None
        assert tbl_w.get(qn("w:w")) == "9360"
        assert tbl_w.get(qn("w:type")) == "dxa"
        assert tbl_ind is not None
        assert tbl_ind.get(qn("w:w")) == "120"
        widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid.gridCol_lst]
        assert sum(widths) == 9360
